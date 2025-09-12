"""Document parser for extracting text from various file formats."""

import logging
from pathlib import Path

try:
    # Prefer the PyPDF2 namespace if available for backward compatibility
    from PyPDF2 import PdfReader  # type: ignore[import-not-found,unused-ignore]
except ModuleNotFoundError:
    try:
        # Newer package name used by many installs
        from pypdf import PdfReader  # type: ignore[import-not-found,unused-ignore]
    except ModuleNotFoundError:
        # Keep a clear name in module namespace; consumers should check and
        # raise a helpful error when attempting to use PDF parsing.
        PdfReader = None

from docx import Document as DocxDocument

from hlpr.models.document import Document, FileFormat

logger = logging.getLogger(__name__)

# Memory management constants
MAX_MEMORY_FILE_SIZE = 50 * 1024 * 1024  # 50MB
STREAMING_CHUNK_SIZE = 1024 * 1024  # 1MB chunks for streaming


class DocumentParser:
    """Parser for extracting text content from document files.

    Supports PDF, DOCX, TXT, and MD file formats using specialized libraries
    for each format with appropriate error handling.
    """

    @staticmethod
    def parse_file(file_path: str | Path) -> str:
        """Parse a document file and extract its text content.

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text content from the document

        Raises:
            ValueError: If file format is unsupported or parsing fails
            FileNotFoundError: If file doesn't exist
            PermissionError: If file cannot be read
        """
        path = Path(file_path)

        if not path.exists():
            msg = f"File not found: {file_path}"
            raise FileNotFoundError(msg)

        if not path.is_file():
            msg = f"Path is not a file: {file_path}"
            raise ValueError(msg)

        # Determine format from file extension
        extension = path.suffix.lower().lstrip(".")
        try:
            file_format = FileFormat(extension)
        except ValueError:
            msg = f"Unsupported file format: {extension}"
            raise ValueError(msg) from None

        # Check file size for memory management
        file_size = path.stat().st_size
        use_streaming = file_size > MAX_MEMORY_FILE_SIZE

        if use_streaming:
            logger.info(
                "File %s is large (%d MB), using streaming parsing",
                file_path,
                file_size // (1024 * 1024),
            )

        # Parse based on format
        try:
            if file_format == FileFormat.PDF:
                return DocumentParser._parse_pdf(path, streaming=use_streaming)
            if file_format == FileFormat.DOCX:
                return DocumentParser._parse_docx(path, streaming=use_streaming)
            if file_format in (FileFormat.TXT, FileFormat.MD):
                return DocumentParser._parse_text(path, streaming=use_streaming)
            msg = f"Unsupported format: {file_format}"
            raise ValueError(msg)  # noqa: TRY301
        except Exception as e:
            logger.exception("Failed to parse %s", file_path)
            msg = f"Failed to parse document: {e}"
            raise ValueError(msg) from e

    @staticmethod
    def _parse_pdf(file_path: Path, *, streaming: bool = False) -> str:
        """Parse PDF file and extract text content.

        Args:
            file_path: Path to PDF file
            streaming: Whether to use streaming parsing for large files

        Returns:
            Extracted text content

        Raises:
            ValueError: If PDF parsing fails or PDF library not available
        """
        if PdfReader is None:
            msg = "PDF parsing not available - install PyPDF2 or pypdf package"
            raise ValueError(msg)

        try:
            if streaming:
                # Use streaming approach for large files
                return DocumentParser._parse_pdf_streaming(file_path)
            # Use standard parsing for smaller files
            reader = PdfReader(file_path)
            text_content = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text)

            if not text_content:
                msg = "No text content found in PDF (may be image-based)"
                raise ValueError(msg)  # noqa: TRY301

            return "\n\n".join(text_content)

        except Exception as e:
            msg = f"Failed to parse PDF: {e}"
            raise ValueError(msg) from e

    @staticmethod
    def _parse_pdf_streaming(file_path: Path) -> str:
        """Parse large PDF files using streaming approach.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content

        Raises:
            ValueError: If PDF parsing fails
        """
        if PdfReader is None:
            msg = "PDF parsing not available - install PyPDF2 or pypdf package"
            raise ValueError(msg)

        try:
            reader = PdfReader(file_path)
            text_content = []

            # Process pages in batches to manage memory
            batch_size = 10  # Process 10 pages at a time
            total_pages = len(reader.pages)

            for start_idx in range(0, total_pages, batch_size):
                end_idx = min(start_idx + batch_size, total_pages)

                # Process batch of pages
                batch_text = []
                for page_idx in range(start_idx, end_idx):
                    page = reader.pages[page_idx]
                    page_text = page.extract_text()
                    if page_text.strip():
                        batch_text.append(page_text)

                if batch_text:
                    text_content.extend(batch_text)

                # Clear the current batch from memory
                del batch_text

                # Log progress for large files
                if total_pages > 50:
                    progress = (end_idx / total_pages) * 100
                    logger.debug(
                        "PDF parsing progress: %.1f%% (%d/%d pages)",
                        progress,
                        end_idx,
                        total_pages,
                    )

            if not text_content:
                msg = "No text content found in PDF (may be image-based)"
                raise ValueError(msg)  # noqa: TRY301

            return "\n\n".join(text_content)

        except Exception as e:
            msg = f"Failed to parse PDF with streaming: {e}"
            raise ValueError(msg) from e

    @staticmethod
    def _parse_docx(file_path: Path, *, streaming: bool = False) -> str:  # noqa: ARG004
        """Parse DOCX file and extract text content.

        Args:
            file_path: Path to DOCX file
            streaming: Whether to use streaming parsing for large files

        Returns:
            Extracted text content

        Raises:
            ValueError: If DOCX parsing fails
        """
        try:
            doc = DocxDocument(file_path)
            text_content = [p.text for p in doc.paragraphs if p.text.strip()]

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    text_content.extend(
                        [cell.text for cell in row.cells if cell.text.strip()],
                    )

            if not text_content:
                msg = "No text content found in DOCX file"
                raise ValueError(msg)  # noqa: TRY301

            return "\n\n".join(text_content)

        except Exception as e:
            msg = f"Failed to parse DOCX: {e}"
            raise ValueError(msg) from e

    @staticmethod
    def _parse_text(file_path: Path, *, streaming: bool = False) -> str:
        """Parse text-based file (TXT, MD) and return content.

        Args:
            file_path: Path to text file
            streaming: Whether to use streaming parsing for large files

        Returns:
            File content as string

        Raises:
            ValueError: If file cannot be read
        """
        try:
            if streaming:
                # Use streaming approach for large files
                return DocumentParser._parse_text_streaming(file_path)
            # Use standard parsing for smaller files
            with file_path.open("r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError as err:
            msg = "File encoding is not UTF-8 compatible"
            raise ValueError(msg) from err
        except Exception as e:
            msg = f"Failed to read text file: {e}"
            raise ValueError(msg) from e
        else:
            if not content.strip():
                msg = "File is empty or contains no readable text"
                raise ValueError(msg)
            return content

    @staticmethod
    def _parse_text_streaming(file_path: Path) -> str:
        """Parse large text files using streaming approach.

        Args:
            file_path: Path to text file

        Returns:
            File content as string

        Raises:
            ValueError: If file cannot be read
        """
        try:
            content_chunks = []

            with file_path.open("r", encoding="utf-8") as f:
                while True:
                    chunk = f.read(STREAMING_CHUNK_SIZE)
                    if not chunk:
                        break
                    content_chunks.append(chunk)

            content = "".join(content_chunks)
        except UnicodeDecodeError as err:
            msg = "File encoding is not UTF-8 compatible"
            raise ValueError(msg) from err
        except Exception as e:
            msg = f"Failed to read text file with streaming: {e}"
            raise ValueError(msg) from e
        else:
            if not content.strip():
                msg = "File is empty or contains no readable text"
                raise ValueError(msg)
            return content

    @staticmethod
    def validate_document(document: Document) -> bool:
        """Validate that a document can be parsed.

        Args:
            document: Document instance to validate

        Returns:
            True if document can be parsed, False otherwise
        """
        try:
            # Try to parse a small portion to validate
            path = Path(document.path)
            can_parse = False

            if document.format == FileFormat.PDF:
                reader = PdfReader(path)
                if len(reader.pages) > 0:
                    # Try to extract text from first page
                    first_page = reader.pages[0]
                    text = first_page.extract_text()
                    can_parse = bool(text.strip())

            elif document.format == FileFormat.DOCX:
                doc = DocxDocument(path)
                can_parse = bool(doc.paragraphs or doc.tables)

            elif document.format in (FileFormat.TXT, FileFormat.MD):
                with path.open("r", encoding="utf-8") as f:
                    content = f.read(1024)  # Read first 1KB
                    can_parse = bool(content.strip())

        except Exception:  # noqa: BLE001
            return False
        else:
            return can_parse

    @staticmethod
    def get_document_info(file_path: str | Path) -> dict:
        """Get basic information about a document without full parsing.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with document information

        Raises:
            ValueError: If file cannot be analyzed
        """
        path = Path(file_path)

        if not path.exists():
            msg = f"File not found: {file_path}"
            raise FileNotFoundError(msg)

        extension = path.suffix.lower().lstrip(".")
        try:
            file_format = FileFormat(extension)
        except ValueError:
            msg = f"Unsupported file format: {extension}"
            raise ValueError(msg) from None

        info = {
            "path": str(path.absolute()),
            "format": file_format.value,
            "size_bytes": path.stat().st_size,
            "can_parse": False,
        }

        try:
            # Quick validation
            if file_format == FileFormat.PDF:
                reader = PdfReader(path)
                info["page_count"] = len(reader.pages)
                info["can_parse"] = len(reader.pages) > 0

            elif file_format == FileFormat.DOCX:
                doc = DocxDocument(path)
                info["paragraph_count"] = len(doc.paragraphs)
                info["table_count"] = len(doc.tables)
                info["can_parse"] = bool(doc.paragraphs or doc.tables)

            elif file_format in (FileFormat.TXT, FileFormat.MD):
                with path.open("r", encoding="utf-8") as f:
                    sample = f.read(1024)
                    info["sample_text"] = (
                        sample[:200] + "..." if len(sample) > 200 else sample
                    )
                    info["can_parse"] = bool(sample.strip())

        except Exception as e:  # noqa: BLE001
            info["error"] = str(e)

        return info
