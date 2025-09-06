"""Document loading utilities for different file formats."""

from typing import Optional, Dict, Any
from pathlib import Path
import logging
import os
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentLoader:
    """Utility class for loading documents from various file formats."""

    @staticmethod
    def load_file(file_path: str) -> Optional[str]:
        """Load text content from a file based on its extension.

        Args:
            file_path: Path to the file to load

        Returns:
            Extracted text content or None if loading fails
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File does not exist: {file_path}")
            return None

        file_extension = file_path.suffix.lower()

        try:
            if file_extension == '.txt':
                return DocumentLoader._load_txt(file_path)
            elif file_extension == '.pdf':
                return DocumentLoader._load_pdf(file_path)
            elif file_extension == '.docx':
                return DocumentLoader._load_docx(file_path)
            elif file_extension == '.md':
                return DocumentLoader._load_txt(file_path)  # Markdown is plain text
            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                return None
        except Exception as e:
            logger.error(f"Error loading file {file_path}: {e}")
            return None

    @staticmethod
    def _load_txt(file_path: Path) -> Optional[str]:
        """Load text from a plain text file with encoding detection."""
        try:
            # Try UTF-8 first (most common)
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to system default encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return None
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return None

    @staticmethod
    def _load_pdf(file_path: Path) -> Optional[str]:
        """Load text from a PDF file using PyMuPDF (pymupdf).

        PyMuPDF is faster and generally produces better text extraction than
        PyPDF2. We import it lazily and return None with a clear log message
        if it's not installed.
        """
        try:
            import pymupdf as fitz
        except ImportError:
            logger.error("PyMuPDF (pymupdf) is required for PDF processing. Install with: pip install pymupdf")
            return None

        text_parts = []
        doc = None
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text.strip():  # Only add non-empty pages
                    text_parts.append(page_text)
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {e}")
            return None
        finally:
            if doc is not None:
                doc.close()

        return "\n".join(text_parts).strip()

    @staticmethod
    def _load_docx(file_path: Path) -> Optional[str]:
        """Load text from a DOCX file."""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx is required for DOCX processing. Install with: pip install python-docx")
            return None

        try:
            doc = Document(file_path)
            text_parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts).strip() if text_parts else ""
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {e}")
            return None

    @staticmethod
    def get_file_metadata(file_path: str) -> Dict[str, Any]:
        """Get comprehensive metadata about a file.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary with file metadata including timestamps as datetime objects
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return {}

        try:
            stat = file_path.stat()
            return {
                'file_name': file_path.name,
                'file_size': stat.st_size,
                'file_type': file_path.suffix.lower(),
                'modified_time': datetime.fromtimestamp(stat.st_mtime),
                'created_time': datetime.fromtimestamp(stat.st_ctime),
                'accessed_time': datetime.fromtimestamp(stat.st_atime),
                'is_readable': file_path.is_file() and os.access(file_path, os.R_OK),
            }
        except Exception as e:
            logger.error(f"Error getting metadata for {file_path}: {e}")
            return {}
    
    @staticmethod
    def get_supported_extensions() -> set[str]:
        """Get set of supported file extensions.
        
        Returns:
            Set of supported file extensions (including the dot)
        """
        return {'.txt', '.pdf', '.docx', '.md'}
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """Check if a file type is supported.
        
        Args:
            file_path: Path to check
            
        Returns:
            True if the file extension is supported
        """
        return Path(file_path).suffix.lower() in DocumentLoader.get_supported_extensions()